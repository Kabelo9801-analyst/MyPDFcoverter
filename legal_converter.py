
#!/usr/bin/env python3
"""
legal_converter.py — Convert legal documents into:
  1) Plain-English (faithful, context-preserving rewrite)
  2) Bullet-point summary (per section, preserving legal meaning)

Supports batch processing for PDF/DOCX/TXT. Sections are detected from
headings where possible (DOCX) and with regex heuristics (PDF/TXT).

USAGE
-----
1) Install dependencies (Python 3.10+):
   pip install -U openai python-docx pdfplumber PyPDF2 reportlab tqdm

2) Set your API key **securely** (do NOT hardcode keys in code):
   # macOS / Linux
   export OPENAI_API_KEY="sk-..."
   # Windows (PowerShell)
   setx OPENAI_API_KEY "sk-..."

3) Run:
   python legal_converter.py --input ./in --output ./out --model gpt-4o-mini --formats docx,txt

   # For multiple workers (process multiple files in parallel)
   python legal_converter.py --input ./in --output ./out --workers 3

SECURITY NOTE
-------------
Never paste API keys into source files or share them. Use environment variables
like OPENAI_API_KEY. This script intentionally reads the key from the environment.
"""

from __future__ import annotations

import argparse
import concurrent.futures as futures
import dataclasses
import json
import logging
import os
import re
import sys
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple

# Third-party
from openai import OpenAI
import pdfplumber
from docx import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from tqdm import tqdm

# Optional (only used for simple PDF writing if installed)
try:
    from reportlab.pdfgen import canvas  # type: ignore
    from reportlab.lib.pagesizes import LETTER  # type: ignore
    REPORTLAB_AVAILABLE = True
except Exception:
    REPORTLAB_AVAILABLE = False


# --------------------------- Logging ----------------------------------

def setup_logging(out_dir: Path) -> None:
    out_dir.mkdir(parents=True, exist_ok=True)
    log_path = out_dir / "processing.log"
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
        handlers=[
            logging.FileHandler(log_path, encoding="utf-8"),
            logging.StreamHandler(sys.stdout),
        ],
    )
    logging.info("Logging initialized. Log file: %s", log_path)


# --------------------------- Data types --------------------------------

@dataclass
class Section:
    """A logical document section."""
    title: str
    text: str
    number: Optional[str] = None  # e.g., "1.1", "Article II", etc.


@dataclass
class DocBundle:
    """Holds the original file and its parsed sections plus metadata."""
    path: Path
    title: str
    date: Optional[str]
    sections: List[Section]


# --------------------------- Utilities ---------------------------------

def read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def extract_from_pdf(path: Path) -> str:
    """Extract text from a PDF using pdfplumber. Falls back to PyPDF2 if needed."""
    try:
        with pdfplumber.open(str(path)) as pdf:
            texts = []
            for page in pdf.pages:
                t = page.extract_text() or ""
                texts.append(t)
            return "\n".join(texts)
    except Exception as e:
        logging.exception("pdfplumber failed for %s, error: %s", path, e)
        try:
            import PyPDF2  # local import fallback
            reader = PyPDF2.PdfReader(str(path))
            texts = []
            for page in reader.pages:
                t = page.extract_text() or ""
                texts.append(t)
            return "\n".join(texts)
        except Exception as e2:
            logging.exception("PyPDF2 also failed for %s, error: %s", path, e2)
            return ""


def extract_from_docx(path: Path) -> Tuple[List[Section], Dict[str, str]]:
    """
    Parse DOCX into sections using Heading styles when available.
    Returns (sections, core_properties).
    """
    doc = DocxDocument(str(path))
    props = doc.core_properties
    meta = {
        "title": props.title or "",
        "created": (props.created.isoformat() if props.created else ""),
        "modified": (props.modified.isoformat() if props.modified else ""),
    }

    sections: List[Section] = []
    current_title = None
    current_number = None
    current_paras: List[str] = []

    def flush_section():
        nonlocal current_title, current_number, current_paras
        if current_title is not None or current_paras:
            text = "\n".join(current_paras).strip()
            title = current_title or "Section"
            sections.append(Section(title=title, text=text, number=current_number))
        current_title, current_number, current_paras = None, None, []

    heading_pattern = re.compile(r"^((?:Article|Section)\s+[A-Z0-9IVXLC]+|\d+(?:\.\d+)*)(?:\s*[:.\)]\s*)?", re.I)

    for p in doc.paragraphs:
        p_text = p.text.strip()
        style_name = (p.style.name if p.style is not None else "")

        if style_name.startswith("Heading") or heading_pattern.match(p_text):
            # Start new section
            flush_section()
            # Try to split number from title
            m = heading_pattern.match(p_text)
            if m:
                current_number = m.group(1)
                rest = p_text[m.end():].strip(" :-\u2013")
                current_title = rest or p_text
            else:
                current_title = p_text
        else:
            if p_text:
                current_paras.append(p_text)

    flush_section()
    # If no sections recognized, use entire doc as single section
    if not sections:
        all_text = "\n".join(p.text for p in doc.paragraphs)
        sections = [Section(title="Document", text=all_text.strip())]

    return sections, meta


def regex_section_split(text: str) -> List[Section]:
    """
    Heuristic sectionizer for PDF/TXT.
    Splits on common legal heading patterns.
    """
    # Normalize excessive whitespace
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Pattern examples:
    # "Section 1. Title", "1.1 Subsection", "ARTICLE II. DEFINITIONS", "ARTICLE 1: SCOPE"
    heading_regex = re.compile(
        r"(?im)^(?P<h>(?:ARTICLE|Section)\s+[A-Z0-9IVXLC]+[.:)]?\s+[^\n]+|"
        r"(?:\d+(?:\.\d+){0,3})\s+[A-Z][^\n]+|"
        r"[A-Z][A-Z0-9 ,\-\(\)\/]{3,})\n"
    )

    sections: List[Section] = []
    last_idx = 0
    matches = list(heading_regex.finditer(text))

    if not matches:
        return [Section(title="Document", text=text.strip())]

    for i, m in enumerate(matches):
        start = m.start()
        if i == 0 and start > 0:
            # preface before the first heading
            preface = text[:start].strip()
            if preface:
                sections.append(Section(title="Preface", text=preface))

        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        heading_line = m.group("h").strip()
        body = text[m.end():end].strip()

        # Parse number if present
        number_match = re.match(r"(?i)^(ARTICLE|Section)\s+([A-Z0-9IVXLC]+)", heading_line)
        number = number_match.group(2) if number_match else None

        sections.append(Section(title=heading_line, text=body, number=number))

    return sections


def load_document(path: Path) -> DocBundle:
    """
    Read file and return a bundle with sections and metadata.
    """
    suffix = path.suffix.lower()
    title_guess = path.stem.replace("_", " ")
    date_guess = None

    if suffix == ".docx":
        sections, meta = extract_from_docx(path)
        title = meta.get("title") or title_guess
        date_guess = meta.get("created") or meta.get("modified") or None
    elif suffix == ".pdf":
        text = extract_from_pdf(path)
        sections = regex_section_split(text)
        title = title_guess
    elif suffix == ".txt":
        text = read_text_file(path)
        sections = regex_section_split(text)
        title = title_guess
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    return DocBundle(path=path, title=title, date=date_guess, sections=sections)


# --------------------------- Chunking ----------------------------------

def split_into_chunks(text: str, max_chars: int = 12000) -> List[str]:
    """
    Split a section's text into chunks under max_chars.
    Prefer paragraph boundaries; fall back to sentence-ish boundaries.
    """
    if len(text) <= max_chars:
        return [text]

    parts: List[str] = []
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    running = []

    def flush():
        if running:
            parts.append("\n\n".join(running))
            running.clear()

    for p in paragraphs:
        if sum(len(x) for x in running) + len(p) + 2 <= max_chars:
            running.append(p)
        else:
            if len(p) > max_chars:
                # Split very long paragraph on sentence-ish boundaries
                sentences = re.split(r"(?<=[\.\?\!;:])\s+", p)
                cur = []
                for s in sentences:
                    if sum(len(x) for x in cur) + len(s) + 1 <= max_chars:
                        cur.append(s)
                    else:
                        if cur:
                            parts.append(" ".join(cur))
                            cur = [s]
                        else:
                            # Sentence itself too long; hard split
                            parts.append(s[:max_chars])
                            s_rest = s[max_chars:]
                            while len(s_rest) > max_chars:
                                parts.append(s_rest[:max_chars])
                                s_rest = s_rest[max_chars:]
                            if s_rest:
                                parts.append(s_rest)
                            cur = []
                if cur:
                    parts.append(" ".join(cur))
                flush()
            else:
                flush()
                running.append(p)
    flush()
    return parts


# --------------------------- LLM Calls ---------------------------------

class LLMClient:
    """
    Thin wrapper over the OpenAI Python SDK.
    Uses Chat Completions for broad compatibility.
    """
    def __init__(self, model: str, temperature: float = 0.2):
        self.client = OpenAI()  # reads OPENAI_API_KEY from env
        self.model = model
        self.temperature = temperature

    def _chat(self, system_prompt: str, user_prompt: str, max_tokens: int = 1200) -> str:
        # Basic retry with exponential backoff
        delay = 2.0
        for attempt in range(6):
            try:
                resp = self.client.chat.completions.create(
                    model=self.model,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},
                    ],
                    temperature=self.temperature,
                    max_tokens=max_tokens,
                )
                content = resp.choices[0].message.content or ""
                return content.strip()
            except Exception as e:
                logging.warning("LLM call failed (attempt %d): %s", attempt + 1, e)
                time.sleep(delay)
                delay = min(delay * 2, 30)
        raise RuntimeError("LLM call failed after multiple retries.")

    def rewrite_plain_english(self, *, doc_title: str, section: Section, chunk_text: str,
                              part_idx: int, part_total: int) -> str:
        system_prompt = (
            "You are a senior legal writing assistant. Rewrite legal text into clear, plain English "
            "for non-lawyers while preserving every obligation, right, exception, condition, defined term, "
            "numbering, and cross-reference. Do not remove content. Keep defined terms (e.g., 'Effective Date') "
            "capitalized exactly as in the original."
        )
        user_prompt = (
            f"Document: {doc_title}\n"
            f"Section: {section.number or ''} {section.title}\n"
            f"Part {part_idx} of {part_total}\n\n"
            "Task:\n"
            "- Convert the following legal text into plain English, word for word.\n"
            "- Preserve all content, obligations, conditions, exceptions, definitions, numbers, and dates.\n"
            "- Keep any list structure.\n"
            "- Output only the rewritten text for this part (no extra commentary).\n\n"
            f"Original legal text:\n{chunk_text}"
        )
        return self._chat(system_prompt, user_prompt, max_tokens=1800)

    def summarize_bullets(self, *, doc_title: str, section: Section, chunk_text: str,
                          part_idx: int, part_total: int) -> str:
        system_prompt = (
            "You are a legal summarization assistant. Create concise bullet points that retain legal meaning. "
            "Capture duties, rights, prohibitions, amounts, deadlines, conditions precedent/subsequent, remedies, "
            "and cross-references. Stay neutral; no opinions."
        )
        user_prompt = (
            f"Document: {doc_title}\n"
            f"Section: {section.number or ''} {section.title}\n"
            f"Part {part_idx} of {part_total}\n\n"
            "Task:\n"
            "- Convert the following legal text into a concise bullet-point summary.\n"
            "- Keep all legal context and meaning.\n"
            "- Use simple dash '-' bullets (one level). Use line breaks between bullets.\n"
            "- Output only the bullet points for this part.\n\n"
            f"Original legal text:\n{chunk_text}"
        )
        return self._chat(system_prompt, user_prompt, max_tokens=900)


# --------------------------- Post-processing ---------------------------

def clean_bullets(text: str) -> str:
    # Normalize to "- " bullets
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    cleaned: List[str] = []
    for ln in lines:
        ln = re.sub(r"^[•\-\*\u2022]+\s*", "- ", ln)  # convert symbols to "- "
        if not ln.startswith("- "):
            ln = "- " + ln
        cleaned.append(ln)
    return "\n".join(cleaned)


def quality_check(original: str, transformed: str, kind: str, min_ratio: float) -> Optional[str]:
    """
    Return warning string if seems incomplete/inadequate.
    """
    if not transformed.strip():
        return f"{kind}: Empty output."
    ratio = len(transformed) / max(1, len(original))
    if ratio < min_ratio:
        return f"{kind}: Output appears short relative to input (len ratio {ratio:.2f})."
    bad_markers = ["As an AI", "I cannot", "[...]", "cannot access", "I'm unable"]
    if any(bm.lower() in transformed.lower() for bm in bad_markers):
        return f"{kind}: Output contains likely disclaimer or placeholder."
    return None


# --------------------------- Writers -----------------------------------

def ensure_docx_styles(doc: DocxDocument) -> None:
    styles = doc.styles
    if "Body Text" not in [s.name for s in styles]:
        style = styles.add_style("Body Text", WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(11)
    # List Bullet style exists by default in most templates.


def write_docx(path: Path, title: str, sections: List[Tuple[str, Optional[str], str]]) -> None:
    """
    sections: list of (section_title, section_number, section_body)
    """
    doc = DocxDocument()
    ensure_docx_styles(doc)
    doc.add_heading(title, level=1)
    for sec_title, sec_num, body in sections:
        hdr = f"{(sec_num + ' ') if sec_num else ''}{sec_title}"
        doc.add_heading(hdr.strip(), level=2)
        # Preserve bullet lines
        for line in body.splitlines():
            if line.strip().startswith("- "):
                p = doc.add_paragraph(line.strip()[2:])
                p.style = "List Bullet"
            else:
                p = doc.add_paragraph(line)
                p.style = "Body Text"
    doc.save(str(path))


def write_txt(path: Path, title: str, sections: List[Tuple[str, Optional[str], str]]) -> None:
    parts = [f"# {title}"]
    for sec_title, sec_num, body in sections:
        hdr = f"\n\n## {(sec_num + ' ') if sec_num else ''}{sec_title}\n"
        parts.append(hdr + body)
    path.write_text("".join(parts), encoding="utf-8")


def write_pdf_simple(path: Path, title: str, sections: List[Tuple[str, Optional[str], str]]) -> None:
    """
    Very simple PDF writer using reportlab (optional). For rich layout, consider DOCX -> PDF conversion.
    """
    if not REPORTLAB_AVAILABLE:
        logging.warning("reportlab not installed; skipping PDF output for %s", path.name)
        return

    c = canvas.Canvas(str(path), pagesize=LETTER)
    width, height = LETTER
    margin = 50
    x = margin
    y = height - margin

    def draw_line(text: str, bold=False):
        nonlocal y
        if y < margin + 20:
            c.showPage()
            y = height - margin
        if bold:
            c.setFont("Helvetica-Bold", 12)
        else:
            c.setFont("Helvetica", 10)
        for chunk in re.split(r"(\n)", text):
            if chunk == "\n":
                y -= 14
            else:
                # wrap long lines
                words = chunk.split()
                line = ""
                for w in words:
                    test = (line + " " + w).strip()
                    if c.stringWidth(test, "Helvetica", 10) < width - 2 * margin:
                        line = test
                    else:
                        c.drawString(x, y, line)
                        y -= 14
                        line = w
                if line:
                    c.drawString(x, y, line)
                    y -= 14

    draw_line(title, bold=True)
    y -= 10
    for sec_title, sec_num, body in sections:
        header = f"{(sec_num + ' ') if sec_num else ''}{sec_title}"
        y -= 10
        draw_line(header, bold=True)
        y -= 4
        draw_line(body, bold=False)
    c.save()


# --------------------------- Processing --------------------------------

def process_bundle(bundle: DocBundle, llm: LLMClient, out_dir: Path,
                   formats: Sequence[str]) -> Tuple[Path, Path]:
    """
    For a single document: produce plain-English and bullet-summary outputs.
    Returns paths to (plain_docx_or_txt, summary_docx_or_txt) for convenience.
    """
    plain_sections: List[Tuple[str, Optional[str], str]] = []
    summary_sections: List[Tuple[str, Optional[str], str]] = []

    for sec in tqdm(bundle.sections, desc=f"Sections for {bundle.path.name}"):
        chunks = split_into_chunks(sec.text, max_chars=12000)
        plain_outputs = []
        summary_outputs = []
        for idx, ch in enumerate(chunks, start=1):
            plain = llm.rewrite_plain_english(
                doc_title=bundle.title, section=sec,
                chunk_text=ch, part_idx=idx, part_total=len(chunks)
            )
            warn = quality_check(ch, plain, "Plain-English", min_ratio=0.4)
            if warn:
                logging.warning("[%s] %s (%s part %d/%d)", bundle.path.name, warn, sec.title, idx, len(chunks))
            plain_outputs.append(plain)

            summ = llm.summarize_bullets(
                doc_title=bundle.title, section=sec,
                chunk_text=ch, part_idx=idx, part_total=len(chunks)
            )
            warn2 = quality_check(ch, summ, "Summary", min_ratio=0.05)
            if warn2:
                logging.warning("[%s] %s (%s part %d/%d)", bundle.path.name, warn2, sec.title, idx, len(chunks))
            summary_outputs.append(clean_bullets(summ))

        plain_text_whole = "\n\n".join(plain_outputs).strip()
        summary_text_whole = "\n".join(summary_outputs).strip()
        plain_sections.append((sec.title, sec.number, plain_text_whole))
        summary_sections.append((sec.title, sec.number, summary_text_whole))

    # Prepare output paths
    base = bundle.path.stem
    out_dir.mkdir(parents=True, exist_ok=True)

    # Defaults: DOCX & TXT
    plain_paths = []
    summary_paths = []

    if "docx" in formats:
        plain_docx = out_dir / f"{base}_plainEnglish.docx"
        summary_docx = out_dir / f"{base}_summary.docx"
        write_docx(plain_docx, f"{bundle.title} — Plain English", plain_sections)
        write_docx(summary_docx, f"{bundle.title} — Bullet Summary", summary_sections)
        plain_paths.append(plain_docx)
        summary_paths.append(summary_docx)

    if "txt" in formats:
        plain_txt = out_dir / f"{base}_plainEnglish.txt"
        summary_txt = out_dir / f"{base}_summary.txt"
        write_txt(plain_txt, f"{bundle.title} — Plain English", plain_sections)
        write_txt(summary_txt, f"{bundle.title} — Bullet Summary", summary_sections)
        plain_paths.append(plain_txt)
        summary_paths.append(summary_txt)

    if "pdf" in formats:
        plain_pdf = out_dir / f"{base}_plainEnglish.pdf"
        summary_pdf = out_dir / f"{base}_summary.pdf"
        write_pdf_simple(plain_pdf, f"{bundle.title} — Plain English", plain_sections)
        write_pdf_simple(summary_pdf, f"{bundle.title} — Bullet Summary", summary_sections)
        plain_paths.append(plain_pdf)
        summary_paths.append(summary_pdf)

    return (plain_paths[0] if plain_paths else out_dir, summary_paths[0] if summary_paths else out_dir)


def find_input_files(input_dir: Path) -> List[Path]:
    exts = {".pdf", ".docx", ".txt"}
    files = []
    for p in input_dir.rglob("*"):
        if p.suffix.lower() in exts and p.is_file():
            files.append(p)
    return sorted(files)


def process_single_file(path: Path, args) -> Tuple[Path, Path]:
    out_root = Path(args.output).resolve()
    out_dir = out_root / path.stem
    bundle = load_document(path)
    logging.info("Loaded %s: %d sections (title=%r date=%r)", path.name, len(bundle.sections), bundle.title, bundle.date)

    llm = LLMClient(model=args.model, temperature=args.temperature)
    formats = [f.strip().lower() for f in args.formats.split(",")]
    return process_bundle(bundle, llm, out_dir, formats)


def main(argv: Optional[Sequence[str]] = None) -> int:
    parser = argparse.ArgumentParser(description="Legal document converter: plain English + bullet summaries")
    parser.add_argument("--input", required=True, help="Folder containing PDF/DOCX/TXT files")
    parser.add_argument("--output", required=True, help="Folder to write outputs")
    parser.add_argument("--model", default="gpt-4o-mini", help="LLM model name (e.g., gpt-4o-mini)")
    parser.add_argument("--temperature", type=float, default=0.2)
    parser.add_argument("--formats", default="docx,txt", help="Comma list: docx,txt,pdf")
    parser.add_argument("--workers", type=int, default=1, help="Parallel workers across files")
    args = parser.parse_args(argv)

    input_dir = Path(args.input).resolve()
    if not input_dir.exists():
        print(f"Input folder not found: {input_dir}", file=sys.stderr)
        return 2

    setup_logging(Path(args.output))

    api_key_present = bool(os.getenv("OPENAI_API_KEY"))
    if not api_key_present:
        logging.error("OPENAI_API_KEY is not set. Please export it before running.")
        return 3

    files = find_input_files(input_dir)
    if not files:
        logging.warning("No input files found in %s", input_dir)
        return 0

    logging.info("Discovered %d file(s). Starting processing...", len(files))

    results: List[Tuple[Path, Path]] = []
    if args.workers > 1:
        with futures.ThreadPoolExecutor(max_workers=args.workers) as ex:
            fut_to_path = {ex.submit(process_single_file, p, args): p for p in files}
            for fut in futures.as_completed(fut_to_path):
                p = fut_to_path[fut]
                try:
                    res = fut.result()
                    results.append(res)
                    logging.info("Completed %s", p.name)
                except Exception as e:
                    logging.exception("Failed processing %s: %s", p, e)
    else:
        for p in files:
            try:
                res = process_single_file(p, args)
                results.append(res)
                logging.info("Completed %s", p.name)
            except Exception as e:
                logging.exception("Failed processing %s: %s", p, e)

    logging.info("All done. Outputs in: %s", Path(args.output).resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
